"""
将 CCLM 风格的 markdown 论文转换为 Word docx。

规则：
- 标题层级：# → 标题（Title），## → Heading 1，### → Heading 2，#### → Heading 3
- 粗体 **...** 保留为 bold run
- 表格 | ... | 转为 Word 原生表格（CCLM 风格：细线、居中数字）
- 图表引用 "Table X" / "Figure X" 保留为普通文本
- 链接 [text](url) → 显示 text（url 放括号里，CCLM 惯例）
- 中英混排：中文用宋体，英文/数字用 Times New Roman
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor

BASE = Path(__file__).resolve().parent
SRC = BASE / "论文正文v2_中文_CCLM.md"
DST = BASE / "论文正文v2_中文_CCLM.docx"

CN_FONT = "宋体"        # 正文中文
EN_FONT = "Times New Roman"  # 正文英文/数字
HEADING_CN_FONT = "宋体"
HEADING_EN_FONT = "Times New Roman"
# 图注：Segoe UI Symbol 五号(10.5pt)，以正确显示 ✓ ◐ 等符号
CAPTION_EN_FONT = "Segoe UI Symbol"
CAPTION_CN_FONT = "宋体"
CAPTION_SIZE_PT = 10.5


def set_run_font(run, cn_font, en_font, size_pt, bold=False, italic=False):
    """为 run 设置中英混排字体"""
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)



SUPERSCRIPT_DIGITS = {"¹": "1", "²": "2", "³": "3", "⁴": "4"}

def add_author_line(paragraph, text, size_pt=12):
    """作者行：把 Unicode 上标数字及其间的 , 和 * 转为真正的 Word 上标 run"""
    text = text.replace("**", "")
    i = 0
    buf = ""
    def flush(sup=False, s=""):
        if not s:
            return
        r = paragraph.add_run(s)
        set_run_font(r, CN_FONT, EN_FONT, size_pt, bold=False)
        r.font.superscript = sup
    while i < len(text):
        ch = text[i]
        if ch in SUPERSCRIPT_DIGITS:
            flush(False, buf); buf = ""
            sup = SUPERSCRIPT_DIGITS[ch]
            i += 1
            # 吸收后续的 ,数字 / ,* 序列，一并上标
            while i < len(text):
                if text[i] in SUPERSCRIPT_DIGITS:
                    sup += SUPERSCRIPT_DIGITS[text[i]]; i += 1
                elif text[i] == "," and i + 1 < len(text) and (text[i+1] in SUPERSCRIPT_DIGITS or text[i+1] == "*"):
                    sup += ","; i += 1
                elif text[i] == "*":
                    sup += "*"; i += 1
                else:
                    break
            flush(True, sup)
        else:
            buf += ch; i += 1
    flush(False, buf)

def add_runs_with_bold(paragraph, text, cn_font=CN_FONT, en_font=EN_FONT,
                       size_pt=12, base_bold=False):
    """将包含 **bold** 与 *italic* 的文本分段加到 paragraph

    先匹配 **bold** 再匹配 *italic*，避免粗体标记被当成斜体。
    未成对的单个 * （通讯作者标记、`gpt4o_*` 前缀）保持字面输出。
    """
    # 也处理内联代码 `code` → 保留为 normal text (CCLM style removes backticks)
    text = text.replace("`", "")

    pattern = re.compile(r"\*\*([^*]+)\*\*|\*([^*\n]+)\*")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_run_font(r, cn_font, en_font, size_pt, bold=base_bold)
        if m.group(1) is not None:            # **bold**
            r = paragraph.add_run(m.group(1))
            set_run_font(r, cn_font, en_font, size_pt, bold=True)
        else:                                  # *italic*
            r = paragraph.add_run(m.group(2))
            set_run_font(r, cn_font, en_font, size_pt,
                         bold=base_bold, italic=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_run_font(r, cn_font, en_font, size_pt, bold=base_bold)


def set_cell_borders(cell):
    """CCLM 风格：细线黑边"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for border_name in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    # Remove existing borders if any
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def add_table_from_md(doc, rows):
    """md 表格 → Word 表格"""
    if not rows:
        return
    ncols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_header = (i == 0)
            add_runs_with_bold(p, cell_text, size_pt=9, base_bold=is_header)
            # 单元格内不要段前/段后空白，行距单倍——否则文字被顶到格子上方
            set_paragraph_spacing(p, before=3, after=3, line_spacing=1.15)
            set_cell_borders(cell)

    doc.add_paragraph()  # spacing after table


def parse_md_table(lines, start_idx):
    """从 lines[start_idx] 开始解析连续的 table，返回 (rows, new_idx)"""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        # skip separator row | --- | --- |
        if re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def set_paragraph_spacing(paragraph, before=6, after=6, line_spacing=1.5):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Global style: margin
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # Base style: default Paragraph font
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = Pt(12)
    style.element.get_or_add_rPr()
    # Ensure East Asian font
    rPr = style.element.rPr
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), EN_FONT)
    rFonts.set(qn("w:hAnsi"), EN_FONT)

    i = 0
    pending_fig_note = False
    while i < len(lines):
        line = lines[i].rstrip()

        # 内部备注：`<!--internal-->` 之后的内容一律不进 Word（投稿件不得出现自用注释）
        if line.strip().startswith("<!--internal-->"):
            break

        # Markdown 引用块（`> ...`）在本项目里只用于写自用备注，不输出
        if line.lstrip().startswith(">"):
            i += 1
            continue

        # Horizontal rule → skip, we use page breaks only for sections if needed
        if line.strip() == "---":
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # H1 (paper title)
        if line.startswith("# ") and not line.startswith("## "):
            title_text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_bold(p, title_text,
                               cn_font=HEADING_CN_FONT, en_font=HEADING_EN_FONT,
                               size_pt=15, base_bold=True)
            set_paragraph_spacing(p, before=12, after=12, line_spacing=1.5)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            h = line[3:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs_with_bold(p, h,
                               cn_font=HEADING_CN_FONT, en_font=HEADING_EN_FONT,
                               size_pt=12, base_bold=True)
            set_paragraph_spacing(p, before=14, after=6, line_spacing=1.5)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            h = line[4:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs_with_bold(p, h,
                               cn_font=HEADING_CN_FONT, en_font=HEADING_EN_FONT,
                               size_pt=12, base_bold=True)
            set_paragraph_spacing(p, before=10, after=4, line_spacing=1.5)
            i += 1
            continue

        # H4
        if line.startswith("#### "):
            h = line[5:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs_with_bold(p, h,
                               cn_font=HEADING_CN_FONT, en_font=HEADING_EN_FONT,
                               size_pt=12, base_bold=True)
            set_paragraph_spacing(p, before=8, after=2, line_spacing=1.5)
            i += 1
            continue

        # Image: ![alt](path)  — center, scaled to ~15 cm width
        m_img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if m_img:
            alt, img_path = m_img.group(1), m_img.group(2)
            ip = Path(img_path)
            if not ip.is_absolute():
                ip = BASE / ip
            if ip.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(ip), width=Cm(14))
                set_paragraph_spacing(p, before=10, after=4, line_spacing=1.0)
            else:
                print(f"  ⚠️  image missing: {ip}")
            i += 1
            continue

        # Table (start with |)
        if line.strip().startswith("|"):
            rows, i = parse_md_table(lines, i)
            add_table_from_md(doc, rows)
            continue

        # List item "- ..."
        if line.strip().startswith("- "):
            content = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, content, size_pt=12)
            set_paragraph_spacing(p, before=2, after=2, line_spacing=1.5)
            i += 1
            continue

        # Figure caption title "**Figure N. ...**" — Segoe UI Symbol 五号，居中加粗
        # 图注含 ✓ ◐ 等符号，Times New Roman 无对应字形
        if line.startswith("**Figure "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            add_runs_with_bold(p, line, cn_font=CAPTION_CN_FONT,
                               en_font=CAPTION_EN_FONT, size_pt=CAPTION_SIZE_PT)
            set_paragraph_spacing(p, before=4, after=2, line_spacing=1.5)
            pending_fig_note = True     # 下一段是该图的注释
            i += 1
            continue

        # Figure caption notes（紧跟标题的一段）— 同字体字号，但两端对齐、不加粗
        if pending_fig_note:
            pending_fig_note = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            add_runs_with_bold(p, line, cn_font=CAPTION_CN_FONT,
                               en_font=CAPTION_EN_FONT, size_pt=CAPTION_SIZE_PT)
            set_paragraph_spacing(p, before=2, after=6, line_spacing=1.0)
            i += 1
            continue

        # Table caption title "**Table N. ...**" — 正文字号，居中加粗
        if line.startswith("**Table "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            add_runs_with_bold(p, line, size_pt=12)
            set_paragraph_spacing(p, before=6, after=2, line_spacing=1.5)
            i += 1
            continue

        # Table caption notes — md 中以 <!--tbl-note--> 显式标记（置于表格下方）
        if line.startswith("<!--tbl-note-->"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            add_runs_with_bold(p, line[len("<!--tbl-note-->"):], size_pt=12)
            set_paragraph_spacing(p, before=2, after=6, line_spacing=1.0)
            i += 1
            continue

        # Author line — 上标数字/逗号/星号用真正的 Word superscript
        if line.startswith("**Authors**") or line.startswith("**作者**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Pt(0)
            # 标签（Authors / 作者）单独加粗，其余交给上标解析
            sep = "**: " if "**: " in line else "**："
            lbl, rest = line.split(sep, 1)
            lbl = lbl.replace("**", "")
            r = p.add_run(lbl + (": " if sep == "**: " else "："))
            set_run_font(r, CN_FONT, EN_FONT, 12, bold=True)
            add_author_line(p, rest, size_pt=12)
            set_paragraph_spacing(p, before=4, after=4, line_spacing=1.5)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)  # CCLM no indent
        add_runs_with_bold(p, line, size_pt=12)
        set_paragraph_spacing(p, before=4, after=4, line_spacing=1.5)
        i += 1

    doc.save(DST)
    print(f"✅ 生成 Word 文档：{DST}")
    print(f"   大小：{DST.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
